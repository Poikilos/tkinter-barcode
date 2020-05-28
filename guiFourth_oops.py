import threading
import time
import barcode
import img2pdf
from barcode.writer import ImageWriter
import time
from fpdf import FPDF
from PIL import Image
import docx
from docx.api import Document
import os


import sys
sys.setrecursionlimit(5000)

block_cipher = None
import tkinter as tk
from tkinter import ttk,Entry

from tkinter.ttk import Progressbar

from threading import Thread
class ThreadingExample(object):
    """ Threading example class
    The run() method will be started and it will run in the background
    until the application exits.
    """

    def action(self):
        while True:
            print("hello")
            if self.stop==1:break
        self.progress=Progressbar(self.win,orient="horizontal",length=300,mode='determinate')
        self.progress.grid(row=6)
        self.progress["maximum"]=100
        self.start_time=time.time()
        #Starting progress bar here
        self.progress.start()

        self.directoryName="singleFolder"
        
        
        self.folderAddress=os.getcwd()
        self.progress['value']=time.time()-self.start_time
        self.progress.update()
        def makeFolder(self):
            try:
                os.mkdir(self.folderAddress+"\\"+self.directoryName)
            except FileExistsError:
                pass
        self.makeFolder=makeFolder(self)
            
        self.Code_128="code128"
        self.letter_words_1=self.gender_var.get()
        self.letter_words_2=self.gender_var_2.get()
        self.progress['value']=time.time()-self.start_time
        self.progress.update()
        def fasterSolution(self,b):
            self.progress['value']=time.time()-self.start_time
            self.progress.update()
            self.document = Document()
            self.COLUMNS=4
            self.ROWS=1000
            self.table=self.document.add_table(rows=self.ROWS,cols=self.COLUMNS)
            self.table_cells=self.table._cells
            for i in range(self.ROWS):
                self.row_cells=self.table_cells[i*self.COLUMNS:(i+1)*self.COLUMNS]
                for cell in self.row_cells:
                    self.paragraph=cell.paragraphs[0]
                    self.run=self.paragraph.add_run()
                    self.run.add_picture(self.folderAddress+"\\"+self.b+".png",width=350000*0.71,height=350000*0.49)
            self.progress['value']=time.time()-self.start_time
            self.progress.update()
            self.document.save(self.folderAddress+"\\"+"singleFolder"+"\\"+self.b+"_"+str(1)+".docx")
            self.progress['value']=time.time()-self.start_time
            self.progress.update()

         
        """def putIntoDocumentFiles(self,b):
            self.document = Document()
            self.table=document.add_table(rows=100,cols=4)#ROWS=25 For 100 barcodes 
            for row in self.table.rows:
                for cell in row.cells:
                    self.paragraph=cell.paragraphs[0]
                    self.run=self.paragraph.add_run()
                    self.run.add_picture(self.folderAddress+"\\"+self.b+".png",width=350000*0.71,height=350000*0.49)#(width,height)=>dimensions(singleTableRow,singleTableColumn) for singlePage
            self.document.save(self.folderAddress+"\\singleFolder"+"\\"+self.b+"_"+"1"+".docx")
            #for i in range(2):
                #document.save(folderAddress+"\\singleFolder"+"\\"+pic+"_"+str(i+1)+".docx")
        """
        
                
        #print(f'{input_1} and {input_2}')
        self.pageNumber=int(self.page_var.get())
        self.array,self.array2=[],[]
        #print('here')
        for i in range(1,1+self.pageNumber):
            self.progress['value']=time.time()-self.start_time
            self.progress.update()
            self.barcode_Name=self.letter_words_1+self.letter_words_2+"00"+str(i)
            self.ean=barcode.get(self.Code_128,self.barcode_Name,writer=ImageWriter())
            print(self.folderAddress+"\\"+self.barcode_Name)
            self.filename=self.ean.save(self.folderAddress+"\\"+self.barcode_Name)#saves to [MG001.png,MG002.png...MG00(pageNumber).png]
            self.array.append(self.folderAddress+"\\"+self.barcode_Name+".png")
            self.array2.append(self.barcode_Name)
            
        makeFolder(self)
        for a,b in zip(self.array,self.array2):
            self.b=b
            self.progress['value']=time.time()-self.start_time
            self.progress.update()
            
            #self.putIntoDocumentFiles=putIntoDocumentFiles(self,b)
            #self.putIntoDocumentFiles(self,b)
            
            self.fasterSolution=fasterSolution(self,self.b)
            fasterSolution(self,self.b)
        self.progress['value']=time.time()-self.start_time
        self.progress.update()
        #Stopping progress bar here
        self.progress.stop()
        
    def start_thread(self):
        global stop
        self.stop=0
        t=Thread(target=self.action)
        t.start()
        
    def stopped(self):
        global stop
        self.stop=1
        
    def __init__(self, interval=1):
        """ Constructor
        :type interval: int
        :param interval: Check interval, in seconds
        """
        self.stop=0
        self.interval = interval
        self.win=tk.Tk()
        self.win.geometry('350x350')
        self.win.title('GUI')
        #self.start_thread=start_thread
        #self.stopped=stopped

        self.gender_label=ttk.Label(self.win,text='Select input 1 : ')
        self.gender_label.grid(row=0,column=0,sticky=tk.W)

        self.gender_var=tk.StringVar()
        self.gender_combobox=ttk.Combobox(self.win,width=14,textvariable=self.gender_var,state='readonly')
        self.gender_combobox['values']=('A','B','C','D','E','F','G','H','I',
                                   'J','K','L','M','N','O','P','Q','R','S','T','U','W',
                                   'X','Y','Z')
        self.gender_combobox.current(0)
        self.gender_combobox.grid(row=0,column=1)

        self.gender_label_2=ttk.Label(self.win,text='Select input 2 : ')
        self.gender_label_2.grid(row=1,column=0,sticky=tk.W)

        self.gender_var_2=tk.StringVar()
        self.gender_combobox_2=ttk.Combobox(self.win,width=14,textvariable=self.gender_var_2,state='readonly')
        self.gender_combobox_2['values']=('A','B','C','D','E','F','G','H','I',
                                   'J','K','L','M','N','O','P','Q','R','S','T','U','W',
                                   'X','Y','Z')
        self.gender_combobox_2.current(0)
        self.gender_combobox_2.grid(row=1,column=1)

        self.page_label=ttk.Label(self.win,text='Until page number starting from 1: ')
        self.page_label.grid(row=2,column=0,sticky=tk.W)

        self.page_var=tk.IntVar()
        self.page_number=Entry(bd=5,textvariable=self.page_var)
        self.page_number.grid(row=2,column=1,sticky=tk.W)

        #self.thread = threading.Thread(name="action",target=self.action)
        #self.thread.daemon = False                         # Daemonize thread
        #thread.start()                                  # Start the execution
        
        self.submit_button=ttk.Button(self.win,text='print Barcodes',command=self.start_thread)
        self.submit_button.grid(row=4,column=0)
        self.stop_button=ttk.Button(self.win,text='Stop',command=self.stopped)
        self.stop_button.grid(row=5,column=0)
        self.win.mainloop()    
        

        
    #def run(self):
       # """ Method that runs forever """
        #while True:
            # Do something
            #print('Doing something imporant in the background')

            #time.sleep(self.interval)
    

if __name__=="__main__":
    ThreadingExample()
#example = ThreadingExample()
#time.sleep(3)
#print('Checkpoint')
#time.sleep(2)
#print('Bye')
