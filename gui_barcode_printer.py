import barcode
import img2pdf
from barcode.writer import ImageWriter
import time
from fpdf import FPDF
from PIL import Image
import docx
from docx.api import Document
import os
import time

#import sys
#sys.setrecursionlimit(5000)

block_cipher = None
import tkinter as tk
from tkinter import ttk,Entry

from tkinter.ttk import Progressbar

win=tk.Tk()
win.title('GUI')


gender_label=ttk.Label(win,text='Select input 1 : ')
gender_label.grid(row=0,column=0,sticky=tk.W)

gender_var=tk.StringVar()
gender_combobox=ttk.Combobox(win,width=14,textvariable=gender_var,state='readonly')
gender_combobox['values']=('A','B','C','D','E','F','G','H','I',
                           'J','K','L','M','N','O','P','Q','R','S','T','U','W',
                           'X','Y','Z')
gender_combobox.current(0)
gender_combobox.grid(row=0,column=1)

gender_label_2=ttk.Label(win,text='Select input 2 : ')
gender_label_2.grid(row=1,column=0,sticky=tk.W)

gender_var_2=tk.StringVar()
gender_combobox_2=ttk.Combobox(win,width=14,textvariable=gender_var_2,state='readonly')
gender_combobox_2['values']=('A','B','C','D','E','F','G','H','I',
                           'J','K','L','M','N','O','P','Q','R','S','T','U','W',
                           'X','Y','Z')
gender_combobox_2.current(0)
gender_combobox_2.grid(row=1,column=1)

page_label=ttk.Label(win,text='Until page number starting from 1: ')
page_label.grid(row=2,column=0,sticky=tk.W)

page_var=tk.StringVar()
page_number=Entry(bd=5,textvariable=page_var)
page_number.grid(row=2,column=1,sticky=tk.W)

def action():
    progress=Progressbar(win,orient="horizontal",length=300,mode='determinate')
    progress.grid(row=5)
    progress["maximum"]=100
    start_time=time.time()
    #Starting progress bar here
    progress.start()
    
    folderAddress=os.getcwd()
    progress['value']=time.time()-start_time
    progress.update()
    def makeFolder(folderAddress,directoryName):
        try:
            os.mkdir(folderAddress+"\\"+directoryName)
        except FileExistsError:
            pass
        
    Code_128="code128"
    letter_words_1=gender_var.get()
    letter_words_2=gender_var_2.get()
    progress['value']=time.time()-start_time
    progress.update()
    def fasterSolution(folderAddress,pic,progress):
        progress['value']=time.time()-start_time
        progress.update()
        document = Document()
        COLUMNS=4
        ROWS=1000
        table=document.add_table(rows=ROWS,cols=COLUMNS)
        table_cells=table._cells
        for i in range(ROWS):
            row_cells=table_cells[i*COLUMNS:(i+1)*COLUMNS]
            for cell in row_cells:
                paragraph=cell.paragraphs[0]
                run=paragraph.add_run()
                run.add_picture(folderAddress+"\\"+pic+".png",width=350000*0.71,height=350000*0.49)
        progress['value']=time.time()-start_time
        progress.update()
        document.save(folderAddress+"\\"+"singleFolder"+"\\"+pic+"_"+str(1)+".docx")
        progress['value']=time.time()-start_time
        progress.update()
        
    def putIntoDocumentFiles(folderAddress,pic):
        document = Document()
        table=document.add_table(rows=100,cols=4)#ROWS=25 For 100 barcodes 
        for row in table.rows:
            for cell in row.cells:
                paragraph=cell.paragraphs[0]
                run=paragraph.add_run()
                run.add_picture(folderAddress+"\\"+pic+".png",width=350000*0.71,height=350000*0.49)#(width,height)=>dimensions(singleTableRow,singleTableColumn) for singlePage
        document.save(folderAddress+"\\singleFolder"+"\\"+pic+"_"+"1"+".docx")
        #for i in range(2):
            #document.save(folderAddress+"\\singleFolder"+"\\"+pic+"_"+str(i+1)+".docx")
            
    #print(f'{input_1} and {input_2}')
    pageNumber=int(page_var.get())
    array,array2=[],[]
    #print('here')
    for i in range(1,1+pageNumber):
        progress['value']=time.time()-start_time
        progress.update()
        barcode_Name=letter_words_1+letter_words_2+"00"+str(i)
        ean=barcode.get(Code_128,barcode_Name,writer=ImageWriter())
        print(folderAddress+"\\"+barcode_Name)
        filename=ean.save(folderAddress+"\\"+barcode_Name)#saves to [MG001.png,MG002.png...MG00(pageNumber).png]
        array.append(folderAddress+"\\"+barcode_Name+".png")
        array2.append(barcode_Name)
    makeFolder(folderAddress,"singleFolder")
    for a,b in zip(array,array2):
        progress['value']=time.time()-start_time
        progress.update()
        #putIntoDocumentFiles(folderAddress,b)
        fasterSolution(folderAddress,b,progress)
    progress['value']=time.time()-start_time
    progress.update()
    #Stopping progress bar here
    progress.stop()
submit_button=ttk.Button(win,text='print Barcodes',command=action)
submit_button.grid(row=4,column=0)
win.mainloop()
