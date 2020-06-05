import threading
import time
import barcode
import img2pdf
from barcode.writer import ImageWriter
import time
from fpdf import FPDF
from PIL import Image
import docx
from docx.api import Document
import os

import sys
sys.setrecursionlimit(5000)

block_cipher = None

import tkinter as tk

from tkinter import *
import tkinter.filedialog
import tkinter.messagebox

from tkinter import ttk,Entry,Text,Frame
from tkinter.scrolledtext import ScrolledText
from tkinter.ttk import Progressbar,Scrollbar

from threading import Thread

import re

import pytube
from pytube import YouTube
class ThreadingExample(object):
    """ Threading example class
    The run() method will be started and it will run in the background
    until the application exits.
    """

    #function downloads youtube video with audio or silence according to a variable named youtube_audio_var
    #if empty url is inserted or wrong pattern matching url is inserted, throws error to textBox accordingly
    def download_youtube_video_with_audio(self):
        video_url=self.youtube_url_var.get()
        yav=self.youtube_audio_var.get()
        if video_url=='':
            self.textBox.insert(1.0,"Empty URL inserted"+"\n")
        else:
            p=re.compile(r'^(http(s)?:\/\/)?((w){3}.)?youtu(be|.be)?(\.com)?\/.+')
            m=p.match(video_url)
            if m:
                try:
                    if yav==True:
                        stream1=YouTube(video_url).streams.filter(progressive=yav)
                        if stream1==[] or stream1 ==None:
                            self.textBox.insert(1.0,"Choose without audio instead of  progressive(video and audio):"+"\n")
                        else:
                            stream1[0].download(self.folderAddress)
                            self.textBox.insert(1.0,"Video finished downloading: "+video_url+" : "+"with audio :"+str(yav)+"\n")
                    elif yav==False:
                        stream2=YouTube(video_url).streams.filter(adaptive=not yav)
                        if stream2==[] or stream2 ==None:
                            self.textBox.insert(1.0,"Choose with audio instead of  adaptive(video):"+"\n")
                        else:
                            stream2[0].download(self.folderAddress)
                            #YouTube(video_url).streams.first().download(self.folderAddress)
                            self.textBox.insert(1.0,"Video finished downloading: "+video_url+" : "+"without audio :"+str(not yav)+"\n")
                except Exception as e:
                    self.textBox.insert(1.0,str(e)+"\n")
                self.youtube_url.delete(0,'end')
            else:
                
                self.textBox.insert(1.0,"Invalid url for download inserted"+"\n")
        
    #function downloads youtube video between the available resolutions
    #if empty url is inserted or wrong pattern matching url is inserted, throws error to textBox accordingly
    def download_youtube_video(self):
        video_url=self.youtube_url_var.get()
        youtube_res=self.youtube_var.get()
        
        if video_url=='':
            self.textBox.insert(1.0,"Empty URL inserted"+"\n")
        else:
            p=re.compile(r'^(http(s)?:\/\/)?((w){3}.)?youtu(be|.be)?(\.com)?\/.+')
            m=p.match(video_url)
            if m:
                try:
                    stream=YouTube(video_url).streams.get_by_resolution(youtube_res)
                    if stream==[] or stream ==None:
                        self.textBox.insert(1.0,"Choose another resolution instead of  :"+youtube_res+"\n")
                    else:
                        stream.download()
                        #YouTube(video_url).streams.first().download(self.folderAddress)
                        self.textBox.insert(1.0,"Video finished downloading: "+video_url+" : "+youtube_res+"\n")
                except Exception as e:
                    self.textBox.insert(1.0,str(e)+"\n")
                self.youtube_url.delete(0,'end')
            else:
                
                self.textBox.insert(1.0,"Invalid url for download inserted"+"\n")

    #function makes a folder named singleFolder if doesn't exists and put corresponding documents of barcodes
    #code_128 barcodes are taken as input from two label above and then given as input to another function
    #sub functions are defined as putIntoDocumentFiles and a faster version of it 
    def action(self):
      
        def makeFolder(self):
            try:
                os.mkdir(self.folderAddress+"\\"+self.directoryName)
            except FileExistsError:
                pass
        self.makeFolder=makeFolder(self)
            
        self.Code_128="code128"
        self.letter_words_1=self.gender_var.get()
        self.letter_words_2=self.gender_var_2.get()

        def fasterSolution(self,b):
            self.document = Document()
            self.COLUMNS=4
            self.ROWS=100
            self.table=self.document.add_table(rows=self.ROWS,cols=self.COLUMNS)
            self.table_cells=self.table._cells
            for i in range(self.ROWS):
                self.row_cells=self.table_cells[i*self.COLUMNS:(i+1)*self.COLUMNS]
                for cell in self.row_cells:
                    self.paragraph=cell.paragraphs[0]
                    self.run=self.paragraph.add_run()
                    self.run.add_picture(self.folderAddress+"\\"+self.b+".png",width=350000*0.71,height=350000*0.49)
            self.document.save(self.folderAddress+"\\"+"singleFolder"+"\\"+self.b+"_"+str(1)+".docx")
           

        
                
        #print(f'{input_1} and {input_2}')
        self.pageNumber=int(self.page_var.get())
        self.array,self.array2=[],[]
        #print('here')
        for i in range(1,1+self.pageNumber):
            self.barcode_Name=self.letter_words_1+self.letter_words_2+"00"+str(i)
            self.ean=barcode.get(self.Code_128,self.barcode_Name,writer=ImageWriter())
            #------------------------
            #print(self.folderAddress+"\\"+self.barcode_Name)
            self.textBox.insert(1.0,self.folderAddress+"\\"+self.barcode_Name+"\n")
            self.filename=self.ean.save(self.folderAddress+"\\"+self.barcode_Name)#saves to [MG001.png,MG002.png...MG00(pageNumber).png]
            self.array.append(self.folderAddress+"\\"+self.barcode_Name+".png")
            self.array2.append(self.barcode_Name)
            
        makeFolder(self)
        for a,b in zip(self.array,self.array2):
            self.b=b
            #self.putIntoDocumentFiles=putIntoDocumentFiles(self,b)
            #self.putIntoDocumentFiles(self,b)
            self.fasterSolution=fasterSolution(self,self.b)
            fasterSolution(self,self.b)
            self.textBox.insert(1.0,"Done:"+self.folderAddress+"\\"+self.directoryName+"\\"+self.b+"_"+str(1)+".docx"+"\n")

    def create_thread_version_Two(self,function_name):
        self.run_thread = Thread(target=function_name)
        self.run_thread.setDaemon(True) # <=== add this line
        self.run_thread.start()
        self.textBox.insert(1.0,str(self.run_thread)+"\n")
        self.textBox.insert(1.0,str(self.run_thread.isAlive())+"\n")
        
        
    def create_thread(self):
        self.run_thread = Thread(target=self.action)
        self.run_thread.setDaemon(True) # <=== add this line
        self.run_thread.start()
        #-------------------
        #print(self.run_thread)
        #print('createThread():',self.run_thread.isAlive())
        self.textBox.insert(1.0,str(self.run_thread)+"\n")
        self.textBox.insert(1.0,str(self.run_thread.isAlive())+"\n")
        #-------------------

    def create_thread_two(self):
        self.run_thread = Thread(target=self.download_youtube_video)
        self.run_thread.setDaemon(True) # <=== add this line
        self.run_thread.start()
        
        self.textBox.insert(1.0,str(self.run_thread)+"\n")
        self.textBox.insert(1.0,str(self.run_thread.isAlive())+"\n")

    def create_thread_three(self):
        self.run_thread = Thread(target=self.download_youtube_video_with_audio)
        self.run_thread.setDaemon(True) # <=== add this line
        self.run_thread.start()
        
        self.textBox.insert(1.0,str(self.run_thread)+"\n")
        self.textBox.insert(1.0,str(self.run_thread.isAlive())+"\n")

    #function sets title to Untitled and deletes everything from textBox    
    def new_file(self,event=None):
        self.win.title("Untitled")
        global file_name
        file_name = None
        self.textBox.delete(1.0, END)
        #on_content_changed()

    #function opens file from devices, sets title to opened file name and insert text into textBox
    def open_file(self,event=None):
        input_file_name = tkinter.filedialog.askopenfilename(defaultextension=".txt",
                                                             filetypes=[("All Files", "*.*"), ("Text Documents", "*.txt")])
        if input_file_name:
            global file_name
            file_name = input_file_name
            self.win.title('{} - {}'.format(os.path.basename(file_name), 'GUI'))
            self.textBox.delete(1.0, END)
            with open(file_name) as _file:
                self.textBox.insert(1.0, _file.read())
            #on_content_changed()

    #function writes text from textBox into file_name
    def write_to_file(self,file_name):
        try:
            content = self.textBox.get(1.0, 'end')
            with open(file_name, 'w') as the_file:
                the_file.write(content)
        except IOError:
            tkinter.messagebox.showwarning("Save", "Could not save the file.")

  
    #function takes file name as input and then runs subfunctions and also sets title to saved file path
    def save_as(self,event=None):
        input_file_name = tkinter.filedialog.asksaveasfilename(defaultextension=".txt",
                                                               filetypes=[("All Files", "*.*"), ("Text Documents", "*.txt")])
        if input_file_name:
            #global self.file_name
            self.file_name = input_file_name
            self.write_to_file(self.file_name)
            self.win.title('{} - {}'.format(os.path.basename(self.file_name), 'GUI'))
        return "break"

    #function saves file_name if not exists as save as function, else write_to_file function
    def save(self,event=None):
        #global self.file_name
        
        
        if not self.file_name:
            self.save_as()
        else:
            self.write_to_file(self.file_name)
        return "break"


    def exit_editor(self,event=None):
        if tkinter.messagebox.askokcancel("Quit?", "Really quit?"):
            self.win.destroy()

    #function creates menubar with submenus as file menubar and edit menubar correspondingly
    def create_menubar(self):
        new_file_icon = PhotoImage(file='icons//new_file.gif')
        open_file_icon = PhotoImage(file='icons/open_file.gif')
        save_file_icon = PhotoImage(file='icons/save.gif')
        cut_icon = PhotoImage(file='icons/cut.gif')
        copy_icon = PhotoImage(file='icons/copy.gif')
        paste_icon = PhotoImage(file='icons/paste.gif')
        undo_icon = PhotoImage(file='icons/undo.gif')
        redo_icon = PhotoImage(file='icons/redo.gif')

        self.menu_bar=Menu(self.win)
        self.file_menu = Menu(self.menu_bar, tearoff=0)
        self.file_menu.add_command(label='New', accelerator='Ctrl+N', compound='left',
                      image=new_file_icon, underline=0, command=self.new_file)
        self.file_menu.add_command(label='Open', accelerator='Ctrl+O', compound='left',
                      image=open_file_icon, underline=0, command=self.open_file)
        self.file_menu.add_command(label='Save', accelerator='Ctrl+S',
                      compound='left', image=save_file_icon, underline=0, command=self.save)
        self.file_menu.add_command(label='Save as', accelerator='Shift+Ctrl+S', command=self.save_as)
        self.file_menu.add_separator()
        self.file_menu.add_command(label='Exit', accelerator='Alt+F4', command=self.exit_editor)
        self.menu_bar.add_cascade(label='File', menu=self.file_menu)
        
        self.win.config(menu=self.menu_bar)        

    #function create widgets like Label, Combobox, Entry and Button
    def create_widget(self):
        #made a label at zero row and zero column to select input between A-Z at the corresponding combobox
        self.gender_label=ttk.Label(self.win,text='Select input 1 : ')
        self.gender_label.grid(row=0,column=0,sticky=tk.W)

        self.gender_var=tk.StringVar()
        self.gender_combobox=ttk.Combobox(self.win,width=14,textvariable=self.gender_var,state='readonly')
        self.gender_combobox['values']=('A','B','C','D','E','F','G','H','I',
                                   'J','K','L','M','N','O','P','Q','R','S','T','U','W',
                                   'X','Y','Z')
        self.gender_combobox.current(0)
        self.gender_combobox.grid(row=0,column=1)

        #made a label at first row and zero column to select input between A-Z at the corresponding combobox
        self.gender_label_2=ttk.Label(self.win,text='Select input 2 : ')
        self.gender_label_2.grid(row=1,column=0,sticky=tk.W)

        self.gender_var_2=tk.StringVar()
        self.gender_combobox_2=ttk.Combobox(self.win,width=14,textvariable=self.gender_var_2,state='readonly')
        self.gender_combobox_2['values']=('A','B','C','D','E','F','G','H','I',
                                   'J','K','L','M','N','O','P','Q','R','S','T','U','W',
                                   'X','Y','Z')
        self.gender_combobox_2.current(0)
        self.gender_combobox_2.grid(row=1,column=1)

        #made a label at second row and zero column to take corresponding entry as input between 1-inf
        self.page_label=ttk.Label(self.win,text='Until page number starting from 1: ')
        self.page_label.grid(row=2,column=0,sticky=tk.W)

        self.page_var=tk.IntVar()
        self.page_number=Entry(bd=5,textvariable=self.page_var)
        self.page_number.grid(row=2,column=1,sticky=tk.W)

        #self.thread = threading.Thread(name="action",target=self.action)
        #self.thread.daemon = False                         # Daemonize thread
        #thread.start()                                  # Start the execution

        #submit button is attached to function action here which is inside create_thread function
        self.submit_button=ttk.Button(self.win,text='print Barcodes',command=self.create_thread)
        self.submit_button.grid(row=4,column=0)

        #made a label at fifth row and zero column that takes youtube url as input
        self.youtube_label=ttk.Label(self.win,text='Give youtube url : ')
        self.youtube_label.grid(row=5,column=0,sticky=tk.W)

        self.youtube_url_var=tk.StringVar()
        self.youtube_url=Entry(bd=5,textvariable=self.youtube_url_var)
        self.youtube_url.grid(row=5,column=1,sticky=tk.W)

        self.youtube_var=tk.StringVar()
        self.youtube_combobox=ttk.Combobox(self.win,width=14,textvariable=self.youtube_var,state='readonly')
        self.youtube_combobox['values']=("720p","480p","360p","240p","144p")
        self.youtube_combobox.current(0)
        self.youtube_combobox.grid(row=5,column=2)

        #made a label at fifth row and zero column that takes youtube url with or without audio as input
        self.youtube_label_2=ttk.Label(self.win,text='Prefer it with audio or silence : ')
        self.youtube_label_2.grid(row=5,column=3,sticky=tk.W)

        self.youtube_audio_var=tk.BooleanVar()
        self.youtube_combobox_2=ttk.Combobox(self.win,width=14,textvariable=self.youtube_audio_var,state='readonly')
        self.youtube_combobox_2['values']=(True,False)
        self.youtube_combobox_2.current(0)
        self.youtube_combobox_2.grid(row=5,column=4)

        #youtube button is attached to function download_youtube_video here which is inside create_thread_2 function
        self.youtube_button=ttk.Button(self.win,text='download youtube video from a url',command=self.create_thread_two)
        self.youtube_button.grid(row=6,column=0)

        #WRONG WAY OF DOING IT
        #self.youtube_button_1=ttk.Button(self.win,text='download youtube video from a url with audio or not',command=self.create_thread_version_Two(self.download_youtube_video_with_audio))
        self.youtube_button_1=ttk.Button(self.win,text='download youtube video from a url with audio or not',command=self.create_thread_three)
        self.youtube_button_1.grid(row=6,column=2)

        '''
        self.frame = Frame(self.win)
        self.frame.grid(column=0,row=6, columnspan=6, rowspan=1, sticky='W')
        self.entry3 = Text(self.frame,height=18)
        self.entry3.pack(side='left', fill='both', expand=True)
        self.scrollbar = Scrollbar(self.frame) # height= not permitted here!
        self.entry3.config(yscrollcommand= self.scrollbar.set)
        self.scrollbar.config(command= self.entry3.yview)
        self.scrollbar.pack(side='right', fill='y')
        #self.stop_button=ttk.Button(self.win,text='Stop',command=self.stopped)
        #self.stop_button.grid(row=5,column=0)
        '''
        self.textBox = ScrolledText(self.win, borderwidth=3, relief="sunken")
        self.textBox.grid(column=0,row=7, columnspan=6, rowspan=1, sticky='W')
        
    def __init__(self, interval=1):
            """ Constructor
            :type interval: int
            :param interval: Check interval, in seconds
            """
            self.stop=0
            self.interval = interval
            self.win=tk.Tk()
            self.win.geometry('800x500')
            self.win.title('GUI')
            self.directoryName="singleFolder"
            self.folderAddress=os.getcwd()
            #self.start_thread=start_thread
            #self.stopped=stopped
            self.create_widget()
            self.file_name=None
            self.create_menubar()
            self.win.mainloop()
            
        
    
        
    #def run(self):
       # """ Method that runs forever """
        #while True:
            # Do something
            #print('Doing something imporant in the background')

            #time.sleep(self.interval)
    

if __name__=="__main__":
    ThreadingExample()
#example = ThreadingExample()
#time.sleep(3)
#print('Checkpoint')
#time.sleep(2)
#print('Bye')
